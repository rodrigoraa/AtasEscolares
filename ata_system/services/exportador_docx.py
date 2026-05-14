from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from services.assinaturas import desserializar_assinaturas


EXPORT_DIR = Path(__file__).resolve().parent.parent / "exports"


def exportar_ata_docx(ata) -> Path:
    EXPORT_DIR.mkdir(exist_ok=True)
    caminho = EXPORT_DIR / f"ata_{ata.numero}_{ata.ano}_{ata.id}.docx"

    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(12)

    titulo = document.add_heading(f"Ata nº {ata.numero}/{ata.ano}", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for bloco in ata.texto_gerado.split("\n\n"):
        paragrafo = document.add_paragraph(bloco.strip())
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragrafo.paragraph_format.first_line_indent = Pt(28)
        paragrafo.paragraph_format.space_after = Pt(8)

    for assinatura in desserializar_assinaturas(ata.assinaturas):
        document.add_paragraph()
        linha = document.add_paragraph("____________________________________________")
        linha.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rotulo = document.add_paragraph(assinatura)
        rotulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.save(caminho)
    return caminho


def exportar_ata_pdf(ata) -> Path:
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    EXPORT_DIR.mkdir(exist_ok=True)
    caminho = EXPORT_DIR / f"ata_{ata.numero}_{ata.ano}_{ata.id}.pdf"

    document = SimpleDocTemplate(
        str(caminho),
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )
    styles = getSampleStyleSheet()
    titulo_style = ParagraphStyle(
        "AtaTitulo",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=14,
        alignment=TA_CENTER,
        spaceAfter=18,
    )
    texto_style = ParagraphStyle(
        "AtaTexto",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=12,
        leading=16,
        alignment=TA_JUSTIFY,
        firstLineIndent=28,
        spaceAfter=8,
    )
    assinatura_linha_style = ParagraphStyle(
        "AssinaturaLinha",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=12,
        alignment=TA_CENTER,
        spaceBefore=18,
    )
    assinatura_rotulo_style = ParagraphStyle(
        "AssinaturaRotulo",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=11,
        alignment=TA_CENTER,
        spaceAfter=4,
    )

    elementos = [Paragraph(f"Ata nº {ata.numero}/{ata.ano}", titulo_style)]
    for bloco in ata.texto_gerado.split("\n\n"):
        elementos.append(Paragraph(bloco.strip(), texto_style))

    elementos.append(Spacer(1, 12))
    for assinatura in desserializar_assinaturas(ata.assinaturas):
        elementos.append(Paragraph("____________________________________________", assinatura_linha_style))
        elementos.append(Paragraph(assinatura, assinatura_rotulo_style))

    document.build(elementos)
    return caminho
